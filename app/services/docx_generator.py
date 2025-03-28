import io
import base64
import re
import logging
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from docx.table import _Cell, Table
import math

logger = logging.getLogger(__name__)

class DocxGeneratorService:
    def __init__(self):
        # Initialize standard and specialty fonts
        self._register_fonts()
        
        # Initialize CSS parser and style cache
        self.css_styles = {}
        
        # Keep track of document structure
        self.current_section = None
        self.section_stack = []
        
    def _register_fonts(self):
        """Register standard and specialty fonts for use in DOCX."""
        try:
            # Register default fonts if needed
            pass
        except Exception as e:
            logger.warning(f"Error initializing fonts: {str(e)}")

    def generate_docx(self, html_content: str) -> bytes:
        """Convert HTML to DOCX with comprehensive structure and style preservation."""
        try:
            # Create document
            doc = Document()
            
            # Configure default paragraph and document settings
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            style.paragraph_format.space_after = Pt(8)
            
            # Add custom styles
            self._add_custom_styles(doc)
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
            
            # Clean up the HTML content
            html_content = self._clean_html_content(html_content)
            
            # Parse HTML
            # After parsing HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            document_div = soup.find('div', class_='document')
            if document_div:
                logger.debug("Document div found with content length: %d", len(document_div.get_text()))
            else:
                logger.warning("Document div not found in HTML content! Content starts with: %s", 
                            html_content[:100] if html_content else "empty")
            
            # Remove meta and script tags
            for tag in soup.find_all(['meta', 'script']):
                tag.decompose()
            
            # Extract style tags for reference
            styles = soup.find_all('style')
            css_content = "\n".join([style.string for style in styles if style.string])
            
            # Parse CSS content and build style cache
            self._parse_css_styles(css_content)
            
            # Process document structure - analyze document organization
            document_structure = self._analyze_document_structure(soup)
            
            # Process document based on structure
            if document_structure:
                for section_type, element in document_structure:
                    if section_type == 'heading':
                        level = int(element.name[1])
                        heading = doc.add_heading(level=level)
                        self._process_text_content(heading, element)
                    elif section_type == 'article':
                        # Handle article sections properly
                        article_title = element.find(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
                        if article_title:
                            level = int(article_title.name[1])
                            heading = doc.add_heading(level=level)
                            self._process_text_content(heading, article_title)
                        
                        # Process article content (excluding the title)
                        for content in element.children:
                            if isinstance(content, Tag) and content.name not in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                                self._process_element(doc, content)
                    elif section_type == 'table':
                        self._process_table(doc, element)
                    elif section_type == 'text':
                        para = doc.add_paragraph()
                        self._process_text_content(para, element)
                    elif section_type == 'page_break':
                        doc.add_page_break()
                    else:
                        # Process other sections
                        self._process_element(doc, element)
            else:
                # Fallback to basic processing if structure analysis fails
                document_div = soup.find('div', class_='document')
                
                if document_div:
                    # Handle multi-page documents
                    for i, page_div in enumerate(document_div.find_all('div', class_='page')):
                        # Process each page content
                        self._process_content(doc, page_div)
                        
                        # Add page break between pages
                        if i < len(document_div.find_all('div', class_='page')) - 1:
                            doc.add_page_break()
                else:
                    # No document structure found, process entire content
                    self._process_content(doc, soup)
            
            # Add headers and footers if present
            self._add_headers_and_footers(doc, soup)
            
            # Perform final cleanup operations
            self._cleanup_docx(doc)
            
            # Save document to bytes
            docx_stream = io.BytesIO()
            doc.save(docx_stream)
            docx_stream.seek(0)
            
            logger.info("DOCX generation completed successfully")
            return docx_stream.getvalue()
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            # Create a simple error document
            error_doc = Document()
            error_doc.add_heading("Error Creating Document", 0)
            error_doc.add_paragraph(f"An error occurred while creating the document: {str(e)}")
            error_stream = io.BytesIO()
            error_doc.save(error_stream)
            error_stream.seek(0)
            return error_stream.getvalue()
        
    def _parse_css_styles(self, css_content):
        """Parse CSS content and build a style cache for elements."""
        if not css_content:
            return
            
        # Reset style cache
        self.css_styles = {}
        
        try:
            # Simple CSS parser for common selectors
            # Parse class selectors
            class_matches = re.finditer(r'\.([a-zA-Z0-9_-]+)\s*{([^}]+)}', css_content)
            for match in class_matches:
                class_name = match.group(1)
                style_content = match.group(2)
                style_dict = self._parse_style_attributes(style_content)
                self.css_styles[f'.{class_name}'] = style_dict
            
            # Parse element selectors
            element_matches = re.finditer(r'([a-zA-Z0-9_-]+)\s*{([^}]+)}', css_content)
            for match in element_matches:
                element_name = match.group(1)
                if element_name.startswith('.'):  # Skip class selectors already processed
                    continue
                style_content = match.group(2)
                style_dict = self._parse_style_attributes(style_content)
                self.css_styles[element_name] = style_dict
            
            # Parse combined selectors (e.g., ".data-table th")
            combined_matches = re.finditer(r'([a-zA-Z0-9_\-\.\s]+)\s*{([^}]+)}', css_content)
            for match in combined_matches:
                selector = match.group(1).strip()
                if selector.count(' ') > 0:  # Only process combined selectors
                    style_content = match.group(2)
                    style_dict = self._parse_style_attributes(style_content)
                    self.css_styles[selector] = style_dict
                    
        except Exception as e:
            logger.warning(f"Error parsing CSS: {str(e)}")

    def _parse_style_attributes(self, style_content):
        """Parse style attributes into a dictionary."""
        style_dict = {}
        style_attrs = re.finditer(r'([a-zA-Z\-]+)\s*:\s*([^;]+);?', style_content)
        for attr in style_attrs:
            property_name = attr.group(1).strip()
            property_value = attr.group(2).strip()
            style_dict[property_name] = property_value
        return style_dict

    def _get_element_styles(self, element, default_styles=None):
        """Get combined styles for an element from inline and CSS styles."""
        if default_styles is None:
            default_styles = {}
            
        combined_styles = default_styles.copy()
        
        # Apply CSS styles based on element type
        if element.name in self.css_styles:
            element_styles = self.css_styles[element.name]
            combined_styles.update(element_styles)
        
        # Apply CSS styles based on class
        if element.get('class'):
            classes = element.get('class')
            if isinstance(classes, str):
                classes = [classes]
                
            for class_name in classes:
                class_selector = f'.{class_name}'
                if class_selector in self.css_styles:
                    class_styles = self.css_styles[class_selector]
                    combined_styles.update(class_styles)
                
                # Check for combined selectors (e.g., ".data-table th")
                if element.parent:
                    for selector, styles in self.css_styles.items():
                        if ' ' in selector:
                            parts = selector.split(' ')
                            if len(parts) == 2:
                                parent_selector, child_selector = parts
                                parent_match = False
                                child_match = False
                                
                                # Check if parent matches
                                if parent_selector.startswith('.'):
                                    parent_class = parent_selector[1:]
                                    parent_classes = element.parent.get('class', [])
                                    if isinstance(parent_classes, str):
                                        parent_classes = [parent_classes]
                                    parent_match = parent_class in parent_classes
                                else:
                                    parent_match = element.parent.name == parent_selector
                                
                                # Check if child matches
                                if child_selector.startswith('.'):
                                    child_class = child_selector[1:]
                                    child_match = child_class in classes
                                else:
                                    child_match = element.name == child_selector
                                
                                if parent_match and child_match:
                                    combined_styles.update(styles)
        
        # Apply inline styles (highest priority)
        inline_style = element.get('style', '')
        if inline_style:
            style_dict = {}
            style_attrs = re.finditer(r'([a-zA-Z\-]+)\s*:\s*([^;]+);?', inline_style)
            for attr in style_attrs:
                property_name = attr.group(1).strip()
                property_value = attr.group(2).strip()
                style_dict[property_name] = property_value
            combined_styles.update(style_dict)
        
        return combined_styles
    

    def _analyze_document_structure(self, soup):
        """
        Analyze the document structure to maintain proper flow and hierarchy.
        Returns a list of (section_type, element) tuples for ordered processing.
        """
        structure = []
        
        # Get document container if it exists
        document_div = soup.find('div', class_='document')
        
        if document_div:
            # Handle multi-page documents
            pages = document_div.find_all('div', class_='page')
            
            for page_idx, page in enumerate(pages):
                # Extract main components from each page
                page_elements = self._extract_page_structure(page)
                structure.extend(page_elements)
                
                # Add page break if not the last page
                if page_idx < len(pages) - 1:
                    structure.append(('page_break', None))
        else:
            # No document structure, process as single page
            structure = self._extract_page_structure(soup)
        
        return structure
        
    def _extract_page_structure(self, page_soup):
        """Extract the elements from a page in their logical order."""
        elements = []
        
        # Process headings first - they define document structure
        headings = page_soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
        
        # Extract articles and sections
        articles = page_soup.find_all('article')
        sections = page_soup.find_all('section')
        
        # Extract tables
        tables = page_soup.find_all('table')
        
        # Now build the page structure in logical order
        
        # First pass: get all direct children of the page in order
        # This preserves the natural flow of content
        if isinstance(page_soup, Tag):
            for child in page_soup.children:
                if isinstance(child, NavigableString):
                    if child.strip():
                        elements.append(('text', child))
                elif child.name:
                    if child.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                        elements.append(('heading', child))
                    elif child.name == 'article':
                        elements.append(('article', child))
                    elif child.name == 'section':
                        elements.append(('article', child))  # Treat sections similarly to articles
                    elif child.name == 'table':
                        elements.append(('table', child))
                    elif child.name == 'p':
                        elements.append(('text', child))
                    elif child.name in ['div', 'span']:
                        # Check if it contains meaningful content
                        if child.get_text().strip():
                            elements.append(('container', child))
                    elif child.name not in ['style', 'script', 'meta', 'link']:
                        elements.append(('element', child))
        
        return elements

    def _clean_html_content(self, html_content):
        """Clean up HTML content for processing."""
        # If the content is wrapped in quotes (from JSON), remove them
        if html_content and html_content.startswith('"') and html_content.endswith('"'):
            html_content = html_content[1:-1]
        
        # Unescape escaped quotes (from JSON)
        html_content = html_content.replace('\\"', '"')
        
        # Replace escaped newlines with actual newlines
        html_content = html_content.replace('\\n', '\n')
        
        # Remove DOCTYPE declarations, comments, and XML declarations
        html_content = re.sub(r'<!DOCTYPE[^>]*>', '', html_content, flags=re.IGNORECASE)
        html_content = re.sub(r'<!--.*?-->', '', html_content, flags=re.DOTALL)
        html_content = re.sub(r'<\?xml[^>]*\?>', '', html_content)
        html_content = html_content.replace('&nbsp;', ' ')
        
        return html_content

    def _add_custom_styles(self, doc):
        """Add custom styles to the document for better formatting."""
        # Add heading styles
        for i in range(1, 7):
            style_name = f'CustomHeading{i}'
            if style_name not in doc.styles:
                style = doc.styles.add_style(style_name, 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
                style.base_style = doc.styles['Heading' + str(i)]
                style.font.name = 'Arial'
                style.font.bold = True
                style.font.size = Pt(20 - (i * 2))  # Decreasing size for each heading level
                style.paragraph_format.space_before = Pt(12)
                style.paragraph_format.space_after = Pt(6)
        
        # Add styles for special text blocks
        if 'ArticleText' not in doc.styles:
            style = doc.styles.add_style('ArticleText', 1)
            style.base_style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            style.paragraph_format.first_line_indent = Inches(0.25)
            style.paragraph_format.space_after = Pt(8)
        
        # Add styles for RTL text
        if 'RTLParagraph' not in doc.styles:
            style = doc.styles.add_style('RTLParagraph', 1)
            style.base_style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Set RTL text direction using XML
            style_element = style._element
            if hasattr(style_element, 'get_or_add_pPr'):
                pPr = style_element.get_or_add_pPr()
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), "1")
                pPr.append(bidi)
        
        # Add styles for table elements
        if 'TableHeader' not in doc.styles:
            style = doc.styles.add_style('TableHeader', 1)
            style.base_style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            style.font.bold = True
            
        # Add List styles
        if 'CustomBulletList' not in doc.styles:
            style = doc.styles.add_style('CustomBulletList', 1)
            style.base_style = doc.styles['List Bullet']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
        
        if 'CustomNumberList' not in doc.styles:
            style = doc.styles.add_style('CustomNumberList', 1)
            style.base_style = doc.styles['List Number']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            
        # Add specific styles for legal document formatting
        if 'LegalArticle' not in doc.styles:
            style = doc.styles.add_style('LegalArticle', 1)
            style.base_style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.bold = True
            style.font.size = Pt(11)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
            
        # Add centered text style
        if 'CenteredText' not in doc.styles:
            style = doc.styles.add_style('CenteredText', 1)
            style.base_style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    
    def _process_content(self, doc, parent_element):
        """Process all content within the parent element in order."""
        # Find all direct children that need processing
        for child in parent_element.children:
            # Skip empty strings and irrelevant elements
            if isinstance(child, NavigableString):
                if child.strip():
                    para = doc.add_paragraph()
                    para.add_run(self._clean_text(child))
            elif child.name and child.name not in ['meta', 'script', 'style', 'link']:
                self._process_element(doc, child)

    def _process_element(self, doc, element, css_content=None):
        """Process an HTML element based on its type with improved structure handling."""
        if element is None:
            return
            
        if isinstance(element, NavigableString):
            if element.strip():
                para = doc.add_paragraph()
                para.add_run(self._clean_text(element))
            return
            
        # Skip processing certain elements
        if element.name in ['meta', 'script', 'style', 'link']:
            return
        
        # Get consolidated styles for this element
        element_styles = self._get_element_styles(element)
        
        # Check for RTL text direction
        is_rtl = False
        if element_styles.get('direction') == 'rtl':
            is_rtl = True
        
        # Process based on element type
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Process headings with proper level
            level = int(element.name[1])
            heading = doc.add_heading(level=level)
            
            # Apply text alignment from style
            text_align = element_styles.get('text-align')
            if text_align:
                if text_align == 'center':
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif text_align == 'right':
                    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif text_align == 'justify':
                    heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
            self._process_text_content(heading, element, is_rtl=is_rtl)
            
            # Apply custom styling
            style_name = f'CustomHeading{level}'
            if style_name in doc.styles:
                heading.style = style_name
                
        elif element.name == 'p':
            # Process paragraphs
            para = doc.add_paragraph()
            
            # Apply paragraph class-based styling
            if element.get('class'):
                classes = element.get('class')
                if isinstance(classes, str):
                    classes = [classes]
                
                if 'text-content' in classes:
                    para.style = 'ArticleText'
                    
            # Set RTL style if needed
            if is_rtl:
                para.style = 'RTLParagraph'
                
            # Process paragraph content
            self._process_text_content(para, element, is_rtl=is_rtl)
            
            # Apply alignment from style
            text_align = element_styles.get('text-align')
            if text_align:
                if text_align == 'center':
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif text_align == 'right':
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif text_align == 'justify':
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            elif element.name == 'article':
                # Process articles - important for legal documents
                # Articles typically contain heading + paragraphs
                
                # First find heading if present
                article_heading = element.find(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
                if article_heading:
                    level = int(article_heading.name[1])
                    heading = doc.add_heading(level=level)
                    self._process_text_content(heading, article_heading)
                
                # Then process remaining content
                for content in element.children:
                    if content != article_heading:  # Skip the heading we already processed
                        if isinstance(content, NavigableString):
                            if content.strip():
                                para = doc.add_paragraph()
                                para.add_run(self._clean_text(content))
                        elif content.name:
                            self._process_element(doc, content)
        
            elif element.name == 'section':
                # Process sections similar to articles
                section_heading = element.find(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
                if section_heading:
                    level = int(section_heading.name[1])
                    heading = doc.add_heading(level=level)
                    self._process_text_content(heading, section_heading)
                
                # Process section content
                for content in element.children:
                    if content != section_heading:
                        if isinstance(content, NavigableString):
                            if content.strip():
                                para = doc.add_paragraph()
                                para.add_run(self._clean_text(content))
                        elif content.name:
                            self._process_element(doc, content)
                            
            elif element.name == 'table':
                # Process tables with special attention to styles and structure
                self._process_table(doc, element)
                
            elif element.name in ['ul', 'ol']:
                # Process lists with improved nesting
                self._process_list(doc, element, is_rtl=is_rtl)
                
            elif element.name == 'div':
                # Process div containers with attention to class-based styling
                
                # Check if this is a section with a specific class
                class_value = element.get('class', [])
                if isinstance(class_value, str):
                    class_value = [class_value]
                    
                if 'text-content' in class_value:
                    # For text-content sections, process as a paragraph
                    p = doc.add_paragraph()
                    if is_rtl:
                        p.style = 'RTLParagraph'
                    self._process_text_content(p, element, is_rtl=is_rtl)
                    
                    # Apply text alignment from style
                    text_align = element_styles.get('text-align')
                    if text_align:
                        if text_align == 'center':
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif text_align == 'right':
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif text_align == 'justify':
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    # Process child elements
                    self._process_content(doc, element)
                    
            elif element.name == 'br':
                # In standalone context, add a paragraph
                doc.add_paragraph()
                
            elif element.name in ['span', 'strong', 'em', 'b', 'i', 'u', 's', 'strike']:
                # For inline elements that appear at top level, wrap in paragraph
                para = doc.add_paragraph()
                if is_rtl:
                    para.style = 'RTLParagraph'
                self._process_text_content(para, element, is_rtl=is_rtl)
                
                # Apply text alignment from style
                text_align = element_styles.get('text-align')
                if text_align:
                    if text_align == 'center':
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif text_align == 'right':
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif text_align == 'justify':
                        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
            elif element.name == 'a':
                # Handle hyperlinks
                para = doc.add_paragraph()
                if is_rtl:
                    para.style = 'RTLParagraph'
                self._process_text_content(para, element, is_rtl=is_rtl)
                
            elif element.name in ['pre', 'code']:
                # Handle code blocks
                para = doc.add_paragraph()
                run = para.add_run(self._clean_text(element.get_text()))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                
            else:
                # For other elements, try to process children
                if any(child.name for child in element.children if isinstance(child, Tag)):
                    self._process_content(doc, element)
                else:
                    # Element only has text content
                    text = self._clean_text(element.get_text())
                    if text:
                        para = doc.add_paragraph()
                        para.add_run(text)
    
    def _clean_text(self, text):
        """Clean and normalize text content."""
        if text is None:
            return ""
        
        # Replace literal line breaks with spaces
        text = text.replace('\\n', ' ')
        
        # Replace multiple spaces with single space
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()

    def _process_text_content(self, paragraph, element, is_rtl=False):
        """Process the text content of an element including inline formatting."""
        if element is None:
            return
            
        if isinstance(element, NavigableString):
            # Clean text and add as run
            text = self._clean_text(element)
            if text:
                run = paragraph.add_run(text)
                if is_rtl:
                    self._set_run_rtl(run)
            return
        
        # Get element styles
        element_styles = self._get_element_styles(element)
        
        # Check if element has direct text
        if element.string and element.string.strip():
            text = self._clean_text(element.string)
            run = paragraph.add_run(text)
            if is_rtl:
                self._set_run_rtl(run)
            self._apply_text_formatting(run, element)
            return
            
        # Process all children
        for child in element.children:
            if isinstance(child, NavigableString):
                text = self._clean_text(child)
                if text:
                    run = paragraph.add_run(text)
                    if is_rtl:
                        self._set_run_rtl(run)
                    self._apply_text_formatting(run, element)  # Apply parent formatting to text
            elif child.name == 'br':
                # Handle line breaks properly
                run = paragraph.add_run()
                run.add_break(WD_BREAK.LINE)
            elif child.name in ['strong', 'b']:
                # Bold text
                text = self._clean_text(child.get_text())
                run = paragraph.add_run(text)
                run.bold = True
                if is_rtl:
                    self._set_run_rtl(run)
                self._apply_text_formatting(run, child)
            elif child.name in ['em', 'i']:
                # Italic text
                text = self._clean_text(child.get_text())
                run = paragraph.add_run(text)
                run.italic = True
                if is_rtl:
                    self._set_run_rtl(run)
                self._apply_text_formatting(run, child)
            elif child.name == 'u':
                # Underlined text
                text = self._clean_text(child.get_text())
                run = paragraph.add_run(text)
                run.underline = True
                if is_rtl:
                    self._set_run_rtl(run)
                self._apply_text_formatting(run, child)
            elif child.name in ['s', 'strike', 'del']:
                # Strikethrough text
                text = self._clean_text(child.get_text())
                run = paragraph.add_run(text)
                run.font.strike = True
                if is_rtl:
                    self._set_run_rtl(run)
                self._apply_text_formatting(run, child)
            elif child.name == 'a':
                # Hyperlinks
                text = self._clean_text(child.get_text())
                href = child.get('href', '')
                
                if text:
                    # Add hyperlink if possible, or styled text if not
                    try:
                        self._add_hyperlink(paragraph, text, href or "#")
                    except:
                        run = paragraph.add_run(text)
                        run.underline = True
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        if is_rtl:
                            self._set_run_rtl(run)
            elif child.name == 'span':
                # Process spans with style info
                self._process_text_content(paragraph, child, is_rtl)
            else:
                # Process other elements recursively
                self._process_text_content(paragraph, child, is_rtl)

    def _set_run_rtl(self, run):
        """Set RTL text direction for a run."""
        try:
            rPr = run._element.get_or_add_rPr()
            rtl = OxmlElement('w:rtl')
            rtl.set(qn('w:val'), "1")
            rPr.append(rtl)
        except Exception as e:
            logger.warning(f"Failed to set RTL: {str(e)}")
            
    def _add_hyperlink(self, paragraph, text, url):
        """Add a hyperlink to a paragraph."""
        try:
            part = paragraph.part
            r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            
            # Create hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            # Create run
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            # Add style (blue and underlined)
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')
            rPr.append(color)
            
            underline = OxmlElement('w:u')
            underline.set(qn('w:val'), 'single')
            rPr.append(underline)
            
            new_run.append(rPr)
            
            # Add text
            t = OxmlElement('w:t')
            t.text = text
            new_run.append(t)
            
            # Add run to hyperlink
            hyperlink.append(new_run)
            
            # Add hyperlink to paragraph
            paragraph._p.append(hyperlink)
            
            return hyperlink
        except Exception as e:
            logger.warning(f"Failed to add hyperlink: {str(e)}")
            # Fallback to styled text
            run = paragraph.add_run(text)
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
            return run
    
    def _apply_paragraph_alignment(self, paragraph, element):
        """Apply text alignment to a paragraph based on HTML attributes and styles."""
        # Get consolidated styles for this element
        element_styles = self._get_element_styles(element)
        
        # Check for text-align in styles
        text_align = element_styles.get('text-align')
        
        # Check direct alignment attribute as fallback
        align_attr = element.get('align')
        
        # Set alignment based on found values
        if text_align:
            if text_align == 'center':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif text_align == 'right':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif text_align == 'justify':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif text_align == 'left':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align_attr:
            align_value = align_attr.lower()
            if align_value == 'center':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align_value == 'right':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align_value == 'justify':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif align_value == 'left':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _apply_text_formatting(self, run, element):
        """Apply text formatting to a run based on HTML style attributes."""
        # Get consolidated styles for this element
        element_styles = self._get_element_styles(element)
        
        # Apply base formatting from element name
        if element.name in ['strong', 'b']:
            run.bold = True
        if element.name in ['em', 'i']:
            run.italic = True
        if element.name == 'u':
            run.underline = True
        if element.name in ['s', 'strike', 'del']:
            run.font.strike = True
        if element.name == 'sub':
            run.font.subscript = True
        if element.name == 'sup':
            run.font.superscript = True
            
        # Font weight from styles
        font_weight = element_styles.get('font-weight', '')
        if font_weight in ['bold', '700', '800', '900']:
            run.bold = True
        
        # Font style from styles
        font_style = element_styles.get('font-style', '')
        if font_style == 'italic':
            run.italic = True
        
        # Text decoration from styles
        text_decoration = element_styles.get('text-decoration', '')
        if 'underline' in text_decoration:
            run.underline = True
        if 'line-through' in text_decoration:
            run.font.strike = True

        # Font color from styles
        color_value = element_styles.get('color', '')
        if color_value:
            try:
                # Handle hex colors
                if color_value.startswith('#'):
                    hex_color = color_value.lstrip('#')
                    if len(hex_color) == 6:
                        r = int(hex_color[0:2], 16)
                        g = int(hex_color[2:4], 16)
                        b = int(hex_color[4:6], 16)
                        run.font.color.rgb = RGBColor(r, g, b)
                    elif len(hex_color) == 3:  # Handle 3-digit hex colors
                        r = int(hex_color[0] + hex_color[0], 16)
                        g = int(hex_color[1] + hex_color[1], 16)
                        b = int(hex_color[2] + hex_color[2], 16)
                        run.font.color.rgb = RGBColor(r, g, b)
                # Handle rgb() format
                elif color_value.startswith('rgb'):
                    rgb_match = re.search(r'rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', color_value)
                    if rgb_match:
                        r = int(rgb_match.group(1))
                        g = int(rgb_match.group(2))
                        b = int(rgb_match.group(3))
                        run.font.color.rgb = RGBColor(r, g, b)
                # Handle named colors
                elif color_value in ['black', 'white', 'red', 'green', 'blue', 'yellow', 'gray', 'purple', 'orange']:
                    if color_value == 'black':
                        run.font.color.rgb = RGBColor(0, 0, 0)
                    elif color_value == 'white':
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    elif color_value == 'red':
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    elif color_value == 'green':
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif color_value == 'blue':
                        run.font.color.rgb = RGBColor(0, 0, 255)
                    elif color_value == 'yellow':
                        run.font.color.rgb = RGBColor(255, 255, 0)
                    elif color_value == 'gray':
                        run.font.color.rgb = RGBColor(128, 128, 128)
                    elif color_value == 'purple':
                        run.font.color.rgb = RGBColor(128, 0, 128)
                    elif color_value == 'orange':
                        run.font.color.rgb = RGBColor(255, 165, 0)
            except Exception as e:
                logger.warning(f"Failed to set color: {str(e)}")
        
        # Font size from styles
        font_size = element_styles.get('font-size', '')
        if font_size:
            try:
                # Extract numeric part
                size_match = re.search(r'(\d+)(?:px|pt|em|rem)?', font_size)
                if size_match:
                    size_value = int(size_match.group(1))
                    # Convert different units to points
                    if 'px' in font_size:
                        # Convert pixels to points (approximate)
                        size_value = int(size_value * 0.75)
                    elif 'em' in font_size or 'rem' in font_size:
                        # Convert em to points (assuming 1em = 12pt as base)
                        size_value = int(size_value * 12)
                    
                    run.font.size = Pt(size_value)
            except Exception as e:
                logger.warning(f"Failed to set font size: {str(e)}")
        
        # Font family from styles
        font_family = element_styles.get('font-family', '')
        if font_family:
            try:
                # Extract first font in the list and remove quotes
                font_name = re.sub(r'^[\'"]|[\'"]$', '', font_family.split(',')[0].strip())
                run.font.name = font_name
            except Exception as e:
                logger.warning(f"Failed to set font family: {str(e)}")

    def _process_table(self, doc, table_elem, is_rtl=False):
        """
        Process HTML tables with comprehensive structure and formatting preservation.
        Improved to ensure styles, borders, and cell spacing are accurately preserved.
        """
        # Extract header and body rows
        thead = table_elem.find('thead')
        tbody = table_elem.find('tbody')
        tfoot = table_elem.find('tfoot')
        
        # Get all rows in correct order
        rows = []
        header_rows = []
        footer_rows = []
        
        if thead:
            header_rows = thead.find_all('tr')
            rows.extend(header_rows)
        if tbody:
            rows.extend(tbody.find_all('tr'))
        if tfoot:
            footer_rows = tfoot.find_all('tr')
            rows.extend(footer_rows)
        
        # If no explicit structure, get all rows directly
        if not rows:
            rows = table_elem.find_all('tr')
        
        if not rows:
            return
            
        # Analyze table dimensions accounting for rowspan/colspan
        col_counts = []
        for row in rows:
            cells = row.find_all(['td', 'th'])
            col_count = 0
            for cell in cells:
                colspan = int(cell.get('colspan', 1))
                col_count += colspan
            col_counts.append(col_count)
        
        # Determine max columns
        max_cols = max(col_counts) if col_counts else 0
        if max_cols == 0:
            return
        
        # Create table
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        # Helper to create cell map for tracing merged areas
        cell_map = [[None for _ in range(max_cols)] for _ in range(len(rows))]
        
        # Get consolidated styles for this table
        table_styles = self._get_element_styles(table_elem)
        
        # Get class information
        table_class = table_elem.get('class', [])
        if isinstance(table_class, str):
            table_class = [table_class]
        
        # Apply table width if specified in styles
        width_value = table_styles.get('width', '')
        if width_value:
            # Extract numeric part and unit
            width_match = re.search(r'(\d+)([%a-z]+)', width_value)
            if width_match:
                width_val = float(width_match.group(1))
                width_unit = width_match.group(2)
                
                if width_unit == '%':
                    # Apply percentage of page width
                    width_pct = min(100, width_val) / 100
                    
                    # Use proper XML approach to set table width
                    tbl = table._tbl
                    # Check if tblPr exists, create if not
                    tblPr = None
                    for child in tbl:
                        if child.tag.endswith('tblPr'):
                            tblPr = child
                            break
                    
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)
                    
                    # Create or find tblW element
                    tblW = None
                    for child in tblPr:
                        if child.tag.endswith('tblW'):
                            tblW = child
                            break
                    
                    if tblW is None:
                        tblW = OxmlElement('w:tblW')
                        tblPr.append(tblW)
                    
                    # Set width attributes
                    tblW.set(qn('w:w'), str(int(5000 * width_pct)))  # 5000 = 100%
                    tblW.set(qn('w:type'), 'pct')
                elif width_unit in ['px', 'pt']:
                    # Convert pixels/points to twips for absolute width
                    twips_value = width_val
                    if width_unit == 'px':
                        twips_value = width_val * 15  # Approximate conversion
                    elif width_unit == 'pt':
                        twips_value = width_val * 20  # 1pt = 20 twips
                    
                    # Set absolute width
                    tbl = table._tbl
                    tblPr = None
                    for child in tbl:
                        if child.tag.endswith('tblPr'):
                            tblPr = child
                            break
                    
                    if tblPr is None:
                        tblPr = OxmlElement('w:tblPr')
                        tbl.insert(0, tblPr)
                    
                    tblW = None
                    for child in tblPr:
                        if child.tag.endswith('tblW'):
                            tblW = child
                            break
                    
                    if tblW is None:
                        tblW = OxmlElement('w:tblW')
                        tblPr.append(tblW)
                    
                    tblW.set(qn('w:w'), str(int(twips_value)))
                    tblW.set(qn('w:type'), 'dxa')
    
        # Apply table alignment based on multiple style indicators
        margin_left = table_styles.get('margin-left', '')
        margin_right = table_styles.get('margin-right', '')
        text_align = table_styles.get('text-align', '')
        
        # Determine alignment
        alignment = None
        if (margin_left == 'auto' and margin_right == 'auto') or 'data-table' in table_class:
            alignment = 'center'
        elif text_align:
            alignment = text_align
        elif margin_left == 'auto':
            alignment = 'right'
        elif margin_right == 'auto':
            alignment = 'left'
            
        # Apply alignment if determined
        if alignment:
            tbl = table._tbl
            
            # Check if tblPr exists, create if not
            tblPr = None
            for child in tbl:
                if child.tag.endswith('tblPr'):
                    tblPr = child
                    break
            
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)
            
            # Create or find jc element
            jc = None
            for child in tblPr:
                if child.tag.endswith('jc'):
                    jc = child
                    break
            
            if jc is None:
                jc = OxmlElement('w:jc')
                tblPr.append(jc)
            
            # Set alignment value based on determined alignment
            if alignment == 'center':
                jc.set(qn('w:val'), 'center')
            elif alignment == 'right':
                jc.set(qn('w:val'), 'right')
            elif alignment == 'left':
                jc.set(qn('w:val'), 'left')

            # Process each row
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            
            # Track current column position accounting for merged cells
            current_col = 0
            
            # Process cells in this row
            for cell_idx, cell in enumerate(cells):
                # Skip positions already occupied by row-spanning cells from previous rows
                while current_col < max_cols and cell_map[row_idx][current_col] is not None:
                    current_col += 1
                
                # If we've run out of columns, break
                if current_col >= max_cols:
                    break
                
                try:
                    # Get colspan and rowspan
                    colspan = int(cell.get('colspan', 1))
                    rowspan = int(cell.get('rowspan', 1))
                    
                    # Get the cell from the Word table
                    table_cell = table.cell(row_idx, current_col)
                    
                    # Mark this cell and merged region in the cell map
                    for r in range(row_idx, min(row_idx + rowspan, len(rows))):
                        for c in range(current_col, min(current_col + colspan, max_cols)):
                            if r == row_idx and c == current_col:
                                cell_map[r][c] = "ORIGIN"  # This is the top-left cell
                            else:
                                cell_map[r][c] = "MERGED"  # This position is part of a merged cell
                    
                    # Clear existing content in the cell
                    for paragraph in table_cell.paragraphs:
                        if paragraph.runs:
                            for run in paragraph.runs:
                                run._element.getparent().remove(run._element)
                    
                    # Get consolidated styles for this cell
                    cell_styles = self._get_element_styles(cell)
                    
                    # Process cell content - key improvement for handling line breaks
                    cell_content = []
                    
                    # Check if cell has simple text or complex content
                    if len(list(cell.children)) == 1 and isinstance(next(cell.children), NavigableString):
                        # Single text node
                        cell_text = self._clean_text(cell.string)
                        if cell_text:
                            para = table_cell.paragraphs[0] if table_cell.paragraphs else table_cell.add_paragraph()
                            para.add_run(cell_text)
                    else:
                        # Complex content - may contain line breaks, formatting, etc.
                        # First check if content should be split on <br/> tags
                        if cell.find('br'):
                            # Handle <br/> tags by creating multiple paragraphs
                            contents = []
                            current_content = ""
                            
                            for item in cell.contents:
                                if isinstance(item, NavigableString):
                                    current_content += str(item)
                                elif item.name == 'br':
                                    contents.append(current_content.strip())
                                    current_content = ""
                                else:
                                    current_content += str(item)
                            
                            # Add any remaining content
                            if current_content.strip():
                                contents.append(current_content.strip())
                            
                            # Create a paragraph for each content piece
                            for i, content in enumerate(contents):
                                if i == 0 and table_cell.paragraphs:
                                    para = table_cell.paragraphs[0]
                                else:
                                    para = table_cell.add_paragraph()
                                
                                # Handle the possibility of HTML in content
                                try:
                                    soup = BeautifulSoup(f"<div>{content}</div>", 'html.parser')
                                    self._process_text_content(para, soup.div, is_rtl=is_rtl)
                                except:
                                    para.add_run(content)
                        else:
                            # Single paragraph but may have formatting
                            para = table_cell.paragraphs[0] if table_cell.paragraphs else table_cell.add_paragraph()
                            self._process_text_content(para, cell, is_rtl=is_rtl)

                    # Handle cell background color
                    bg_color = cell_styles.get('background-color', '')
                    if bg_color:
                        # Extract hex color
                        if bg_color.startswith('#'):
                            color_hex = bg_color.lstrip('#')
                            if len(color_hex) == 3:  # Convert shorthand hex to full
                                color_hex = color_hex[0]*2 + color_hex[1]*2 + color_hex[2]*2
                            self._set_cell_shading(table_cell, color_hex)
                        elif bg_color.startswith('rgb'):
                            # Handle rgb() format
                            rgb_match = re.search(r'rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)', bg_color)
                            if rgb_match:
                                r = int(rgb_match.group(1))
                                g = int(rgb_match.group(2))
                                b = int(rgb_match.group(3))
                                color_hex = f"{r:02x}{g:02x}{b:02x}"
                                self._set_cell_shading(table_cell, color_hex)
                        elif bg_color in ['lightgray', 'lightgrey', 'gray', 'grey']:
                            # Handle some common named colors
                            if bg_color in ['lightgray', 'lightgrey']:
                                self._set_cell_shading(table_cell, 'd3d3d3')
                            else:
                                self._set_cell_shading(table_cell, '808080')

                    # Special formatting for header cells
                    if cell.name == 'th':
                        # Make header cells bold
                        for paragraph in table_cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                        
                        # Apply header styling (like background color) if not already set
                        if not bg_color and 'TableHeader' in doc.styles:
                            self._set_cell_shading(table_cell, 'f2f2f2')
                    
                    # Apply text alignment from styles
                    text_align = cell_styles.get('text-align', '')
                    align_attr = cell.get('align')
                    
                    if text_align or align_attr:
                        align_value = None
                        if text_align:
                            align_value = text_align.lower()
                        elif align_attr:
                            align_value = align_attr.lower()
                        
                        for paragraph in table_cell.paragraphs:
                            if align_value == 'center':
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif align_value == 'right':
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            elif align_value == 'justify':
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            elif align_value == 'left':
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Apply vertical alignment from styles
                    vertical_align = cell_styles.get('vertical-align', '')
                    valign_attr = cell.get('valign')
                    
                    if vertical_align or valign_attr:
                        valign_value = None
                        if vertical_align:
                            valign_value = vertical_align.lower()
                        elif valign_attr:
                            valign_value = valign_attr.lower()
                        
                        if valign_value == 'top':
                            table_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                        elif valign_value in ['middle', 'center']:
                            table_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        elif valign_value == 'bottom':
                            table_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
                            
                    # Apply borders if specified in styles
                    border_top = cell_styles.get('border-top', '')
                    border_bottom = cell_styles.get('border-bottom', '')
                    border_left = cell_styles.get('border-left', '')
                    border_right = cell_styles.get('border-right', '')
                    border = cell_styles.get('border', '')
                    
                    # Apply individual borders or full border if specified
                    if any([border_top, border_bottom, border_left, border_right, border]):
                        self._set_cell_borders(table_cell, border_top, border_bottom, border_left, border_right, border)
                    
                    # Apply cell spanning
                    if colspan > 1:
                        # Merge cells horizontally
                        for i in range(1, colspan):
                            if current_col + i < max_cols:
                                try:
                                    table_cell.merge(table.cell(row_idx, current_col + i))
                                except Exception as e:
                                    logger.warning(f"Failed to merge cells horizontally: {str(e)}")
                    
                    if rowspan > 1:
                        # Merge cells vertically
                        for i in range(1, rowspan):
                            if row_idx + i < len(rows):
                                try:
                                    target_cell = table.cell(row_idx + i, current_col)
                                    # Only merge if not already part of another merged cell
                                    if cell_map[row_idx + i][current_col] == "MERGED":
                                        table_cell.merge(target_cell)
                                except Exception as e:
                                    logger.warning(f"Failed to merge cells vertically: {str(e)}")
                    
                    # Update current column position
                    current_col += colspan
                    
                except Exception as e:
                    logger.warning(f"Error processing table cell: {str(e)}")
                    current_col += 1
        
        return table

    def _set_cell_shading(self, cell, color_hex):
        """Set the background shading of a table cell."""
        try:
            tcPr = cell._tc.get_or_add_tcPr()
            
            # Remove existing shading if present
            for shd in tcPr.findall(qn('w:shd')):
                tcPr.remove(shd)
                
            # Add new shading
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color_hex)
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            tcPr.append(shading)
        except Exception as e:
            logger.warning(f"Failed to set cell shading: {str(e)}")
            
    def _set_cell_borders(self, cell, border_top, border_bottom, border_left, border_right, border_all):
        """Set borders for a table cell."""
        try:
            tcPr = cell._tc.get_or_add_tcPr()
            
            # Get or create borders element
            tcBorders = None
            for element in tcPr.findall(qn('w:tcBorders')):
                tcBorders = element
                break
                
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
                
            # Function to parse border style
            def parse_border_style(border_str):
                if not border_str:
                    return None, None
                    
                # Default values
                width = 1
                color = '000000'  # Black
                
                # Parse width and color from border string
                width_match = re.search(r'(\d+)px', border_str)
                if width_match:
                    width = int(width_match.group(1))
                    
                color_match = re.search(r'#([0-9a-fA-F]{3,6})', border_str)
                if color_match:
                    hex_color = color_match.group(1)
                    if len(hex_color) == 3:
                        hex_color = hex_color[0]*2 + hex_color[1]*2 + hex_color[2]*2
                    color = hex_color
                    
                # Map width to Word border size
                if width <= 1:
                    size = 'single'
                elif width <= 2:
                    size = 'thick'
                else:
                    size = 'thickThenThick'
                    
                return size, color

    # Apply all borders if specified
            if border_all:
                size, color = parse_border_style(border_all)
                if size and color:
                    # Apply to all sides
                    sides = ['top', 'left', 'bottom', 'right']
                    for side in sides:
                        border_elem = OxmlElement(f'w:{side}')
                        border_elem.set(qn('w:val'), size)
                        border_elem.set(qn('w:color'), color)
                        border_elem.set(qn('w:sz'), '4')  # 4 = 1pt
                        
                        # Replace existing or add new
                        existing = tcBorders.find(qn(f'w:{side}'))
                        if existing is not None:
                            tcBorders.remove(existing)
                        tcBorders.append(border_elem)
            else:
                # Apply individual borders
                borders = [
                    ('top', border_top),
                    ('left', border_left),
                    ('bottom', border_bottom),
                    ('right', border_right)
                ]
                
                for side, border_value in borders:
                    if border_value:
                        size, color = parse_border_style(border_value)
                        if size and color:
                            border_elem = OxmlElement(f'w:{side}')
                            border_elem.set(qn('w:val'), size)
                            border_elem.set(qn('w:color'), color)
                            border_elem.set(qn('w:sz'), '4')  # 4 = 1pt
                            
                            # Replace existing or add new
                            existing = tcBorders.find(qn(f'w:{side}'))
                            if existing is not None:
                                tcBorders.remove(existing)
                            tcBorders.append(border_elem)
                
        except Exception as e:
            logger.warning(f"Failed to set cell borders: {str(e)}")

    def _process_list(self, doc, list_elem, is_rtl=False):
        """
        Process HTML lists with proper nesting and formatting.
        Improved to handle bullet and numbered lists correctly, with better style preservation.
        """
        is_ordered = list_elem.name == 'ol'
        items = list_elem.find_all('li', recursive=False)
        
        # Get list styles
        list_styles = self._get_element_styles(list_elem)
        list_class = list_elem.get('class', [])
        if isinstance(list_class, str):
            list_class = [list_class]
        
        # Determine list level by counting parent lists
        level = 0
        parent = list_elem.parent
        while parent:
            if parent.name in ['ol', 'ul']:
                level += 1
            parent = parent.parent
        
        # Check for list-style-type to determine bullet/number format
        list_style_type = list_styles.get('list-style-type', '')

        # Process list items
        for item in items:
            # Choose appropriate list style based on type and nesting
            if is_ordered:
                if 'CustomNumberList' in doc.styles:
                    style_name = 'CustomNumberList'
                else:
                    style_name = 'List Number'
            else:
                if 'CustomBulletList' in doc.styles:
                    style_name = 'CustomBulletList'
                else:
                    style_name = 'List Bullet'
            
            # Create a paragraph with list style
            p = doc.add_paragraph(style=style_name)
            
            # Apply RTL if needed
            if is_rtl:
                p._p.get_or_add_pPr().append(OxmlElement('w:bidi'))
            
            # Set the list level based on nesting
            if level > 0:
                p._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level
            
            # Get item styles
            item_styles = self._get_element_styles(item)
            
            # Apply alignment if specified
            text_align = item_styles.get('text-align', '')
            if text_align:
                if text_align == 'center':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif text_align == 'right':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif text_align == 'justify':
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Process the content of the list item
            self._process_text_content(p, item, is_rtl=is_rtl)
            
            # Handle nested lists
            nested_lists = item.find_all(['ul', 'ol'], recursive=False)
            for nested_list in nested_lists:
                self._process_list(doc, nested_list, is_rtl=is_rtl)

    def _add_headers_and_footers(self, doc, soup):
        """Add headers and footers to the document if present in HTML."""
        # Find header and footer elements
        header_elem = soup.find('header')
        footer_elem = soup.find('footer')
        
        # Process header
        if header_elem:
            # Get header styles
            header_styles = self._get_element_styles(header_elem)
            
            for section in doc.sections:
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                self._process_text_content(header_para, header_elem)
                
                # Apply alignment based on styles
                text_align = header_styles.get('text-align', '')
                if text_align == 'center':
                    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif text_align == 'left':
                    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    # Default to right alignment for headers
                    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT    

        # Process footer
        if footer_elem:
            # Get footer styles
            footer_styles = self._get_element_styles(footer_elem)
            
            for section in doc.sections:
                footer = section.footer
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                self._process_text_content(footer_para, footer_elem)
                
                # Apply alignment based on styles
                text_align = footer_styles.get('text-align', '')
                if text_align == 'right':
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif text_align == 'left':
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    # Default to center alignment for footers
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                # Add page numbers if specified
                if 'page-number' in footer_styles or footer_elem.get('data-page-numbers'):
                    # Add field code for page numbers
                    run = footer_para.add_run()
                    fldChar1 = OxmlElement('w:fldChar')
                    fldChar1.set(qn('w:fldCharType'), 'begin')
                    run._element.append(fldChar1)
                    
                    instrText = OxmlElement('w:instrText')
                    instrText.set(qn('xml:space'), 'preserve')
                    instrText.text = "PAGE"
                    run._element.append(instrText)
                    
                    fldChar2 = OxmlElement('w:fldChar')
                    fldChar2.set(qn('w:fldCharType'), 'end')
                    run._element.append(fldChar2)

    def get_base64_docx(self, docx_data: bytes) -> str:
        """Convert DOCX data to a base64 string."""
        return base64.b64encode(docx_data).decode('utf-8')
        
    def _cleanup_docx(self, doc):
        """Perform final cleanup operations on the document before saving."""
        # Remove any empty paragraphs at the end of the document
        while doc.paragraphs and not doc.paragraphs[-1].text.strip():
            p = doc.paragraphs[-1]._element
            p.getparent().remove(p)
            
        # Ensure proper spacing and formatting consistency
        for paragraph in doc.paragraphs:
            # Fix potential spacing issues
            if paragraph.style.name.startswith('Heading'):
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(6)
            elif paragraph.style.name == 'Normal':
                paragraph.paragraph_format.space_after = Pt(8)
                
        # Ensure proper section breaks
        for section in doc.sections:
            # Set default page orientation and margins if not explicitly set
            if section.orientation is None:
                section.orientation = 0  # Portrait
            
            # Ensure reasonable margins
            if section.left_margin == 0:
                section.left_margin = Inches(1)
            if section.right_margin == 0:
                section.right_margin = Inches(1)
            if section.top_margin == 0:
                section.top_margin = Inches(1)
            if section.bottom_margin == 0:
                section.bottom_margin = Inches(1)

# Create the singleton instance that can be imported elsewhere
docx_generator_service = DocxGeneratorService()
    