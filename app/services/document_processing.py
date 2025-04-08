import time
import uuid
import os
import io
import gc
import re
import tempfile
import logging
import traceback
import base64
from fastapi import UploadFile
from sqlalchemy.orm import Session
from app.core.config import settings
from app.models.translation import TranslationProgress, TranslationChunk
from typing import Tuple, List, Dict, Any, Optional
import fitz  # PyMuPDF
from bs4 import BeautifulSoup

# Define our own TranslationError class to avoid circular imports
class TranslationError(Exception):
    def __init__(self, message: str, code: str):
        super().__init__(message)
        self.code = code
        self.name = 'TranslationError'

# Configure logger
logger = logging.getLogger("documents")

class DocumentProcessingService:
    @staticmethod
    async def process_file(
         file: UploadFile,
         from_lang: str,
         to_lang: str,
         user_id: str,
         db: Session,
     ) -> Tuple[str, int, str]:
         """Process a file for translation and return the content, page count, and process ID."""
         
         print(f"📢 Processing file: {file.filename} ({file.content_type})")
 
         # Ensure `settings` is correctly referenced
         file_type = file.content_type.lower() if file.content_type else ""
         file_content = await file.read()
         file_size = len(file_content)
 
         # File size validation
         if file_size > settings.MAX_FILE_SIZE:
             print("❌ File too large")
             raise TranslationError(
                 f"File too large. Maximum size is {settings.MAX_FILE_SIZE / (1024 * 1024)}MB.",
                 "VALIDATION_ERROR",
             )
 
         # Supported file type check
         if file_type not in settings.SUPPORTED_IMAGE_TYPES + settings.SUPPORTED_DOC_TYPES:
             print("❌ Unsupported file type")
             raise TranslationError(
                 f"Unsupported file type: {file_type}. Supported types: {', '.join(settings.SUPPORTED_IMAGE_TYPES + settings.SUPPORTED_DOC_TYPES)}",
                 "VALIDATION_ERROR",
             )
 
         # Generate unique process ID
         process_id = str(uuid.uuid4())
         print(f"📄 Starting document processing: {process_id}")
 
         # Create a translation progress record
         translation_progress = TranslationProgress(
             processId=process_id,
             userId=user_id,
             status="in_progress",
             fileName=file.filename,
             fromLang=from_lang,
             toLang=to_lang,
             fileType=file_type,
         )
         db.add(translation_progress)
         db.commit()
 
         try:
             print("✅ Reached processing block")
 
             if file_type in settings.SUPPORTED_IMAGE_TYPES:
                 print("🖼 Processing image file...")
                 # Import translation_service here to avoid circular imports
                 from app.services.translation import translation_service
                 html_content = await translation_service.extract_from_image(file_content)
                 translated_content = await translation_service.translate_chunk(html_content, from_lang, to_lang)
                 content = translated_content
                 total_pages = 1
 
                 translation_progress.totalPages = 1
                 translation_progress.currentPage = 1
                 translation_progress.progress = 100
                 db.add(TranslationChunk(processId=process_id, content=content, pageNumber=1))
                 db.commit()
 
             else:
                 print("📄 Processing PDF file using in-memory approach...")
 
                 # Create a BytesIO object from the file content
                 pdf_buffer = io.BytesIO(file_content)
                 
                 try:
                     # Open PDF directly from memory
                     with fitz.open(stream=pdf_buffer, filetype="pdf") as doc:
                         translated_contents = []
                         total_pages = len(doc)
                         translation_progress.totalPages = total_pages
                         db.commit()
 
                         for page_num in range(total_pages):
                             print(f"📖 Processing page {page_num + 1}/{total_pages}")
                             page = doc[page_num]
 
                             # Update progress
                             translation_progress.currentPage = page_num + 1
                             translation_progress.progress = ((page_num + 1) / total_pages) * 100
                             db.commit()
 
                             # Import translation_service here to avoid circular imports
                             from app.services.translation import translation_service
                             # Extract formatted content using the in-memory version
                             html_content = await translation_service._get_formatted_text_from_gemini_buffer(page)
 
                             if html_content and len(html_content) > 50:
                                 translated_content = await translation_service.translate_chunk(html_content, from_lang, to_lang)
                                 if translated_content:
                                     translated_contents.append(translated_content)
                                     db.add(TranslationChunk(processId=process_id, content=translated_content, pageNumber=page_num + 1))
                                     db.commit()
                                 else:
                                     print(f"⚠️ Translation failed for page {page_num + 1}")
                             else:
                                 print(f"⚠️ No valid content extracted from page {page_num + 1}")
 
                         if not translated_contents:
                             translation_progress.status = "failed"
                             db.commit()
                             raise TranslationError("No content extracted and translated from the document", "CONTENT_ERROR")
 
                         content = translation_service.combine_html_content(translated_contents)
 
                 finally:
                     # Ensure all resources are properly closed
                     if 'pdf_buffer' in locals():
                         pdf_buffer.close()
                     
                     # Force garbage collection
                     gc.collect()
 
             # Final validation of the translation result
             if not content.strip():
                 translation_progress.status = "failed"
                 db.commit()
                 raise TranslationError("Translation resulted in empty content", "CONTENT_ERROR")
 
             if "<" not in content or ">" not in content:
                 translation_progress.status = "failed"
                 db.commit()
                 raise TranslationError("Translation result lacks proper HTML formatting", "CONTENT_ERROR")
 
             translation_progress.status = "completed"
             translation_progress.progress = 100
             db.commit()
 
             print(f"✅ Document processing completed: {process_id}")
             return content, total_pages, process_id
 
         except Exception as e:
             print(f"❌ Error during processing: {e}")
             translation_progress.status = "failed"
             db.commit()
             raise e
    
    async def process_text_document(self, file_content: bytes, file_type: str) -> str:
        """Extract content from text-based documents like DOC, DOCX, ODT, TXT, RTF."""
        try:
            # For text/plain documents
            if file_type == "text/plain":
                # Decode text content directly
                text_content = file_content.decode('utf-8', errors='replace')
                # Convert to simple HTML
                html_content = f"<div class='text-content'>{text_content}</div>"
                return html_content
                
            # For RTF documents
            elif file_type in ["text/rtf", "application/rtf"]:
                # Extract text from RTF - this is simplified
                # In a real implementation, you'd use a library like striprtf
                # For now, we'll just extract text between RTF markers
                text_content = file_content.decode('utf-8', errors='replace')
                # Strip RTF formatting (simplified approach)
                text_content = re.sub(r'\\[a-z]+', ' ', text_content)
                text_content = re.sub(r'[{}]', '', text_content)
                # Convert to simple HTML
                html_content = f"<div class='text-content'>{text_content}</div>"
                return html_content
                
            # For word processor documents (DOC, DOCX, ODT)
            elif file_type in ["application/msword", 
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             "application/vnd.oasis.opendocument.text"]:
                # Import locally to avoid circular imports
                from app.services.translation import translation_service
                
                # Create a temporary file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                    temp_file.write(file_content)
                    temp_file_path = temp_file.name
                    
                try:
                    # Modified approach: Use direct text extraction instead of Gemini for first pass
                    try:
                        # Try to use python-docx for DOCX files
                        import docx
                        doc = docx.Document(temp_file_path)
                        
                        # Extract text from paragraphs
                        paragraphs = []
                        for para in doc.paragraphs:
                            if para.text.strip():
                                paragraphs.append(f"<p>{para.text}</p>")
                        
                        # Extract text from tables
                        tables_html = []
                        for table in doc.tables:
                            table_html = "<table class='data-table'>"
                            for row in table.rows:
                                table_html += "<tr>"
                                for cell in row.cells:
                                    table_html += f"<td>{cell.text}</td>"
                                table_html += "</tr>"
                            table_html += "</table>"
                            tables_html.append(table_html)
                        
                        # Combine content
                        html_content = f"""
                        <style>
                            .document {{
                                width: 100%;
                                max-width: 1000px;
                                margin: 0 auto;
                                font-family: Arial, sans-serif;
                                line-height: 1.5;
                            }}
                            .text-content {{
                                margin-bottom: 1em;
                            }}
                            .data-table {{
                                width: 100%;
                                border-collapse: collapse;
                                margin-bottom: 1em;
                            }}
                            .data-table td, .data-table th {{
                                border: 1px solid black;
                                padding: 0.5em;
                            }}
                        </style>
                        <div class="document">
                            {"".join(paragraphs)}
                            {"".join(tables_html)}
                        </div>
                        """
                        
                        logger.info(f"Successfully extracted content from DOCX using python-docx")
                        return html_content
                        
                    except Exception as docx_err:
                        logger.warning(f"Failed to extract with python-docx: {str(docx_err)}, trying fallback method")
                        
                        # Fallback method - convert DOCX to PDF first, since Gemini supports PDF
                        try:
                            # Try to convert DOCX to PDF using a library like docx2pdf
                            import subprocess
                            
                            # Create a temporary output PDF file
                            pdf_output_path = temp_file_path + ".pdf"
                            
                            # Try using soffice (LibreOffice) for conversion if available
                            try:
                                # Check if soffice is available
                                subprocess.run(["which", "soffice"], check=True, capture_output=True)
                                
                                # Use LibreOffice for conversion
                                cmd = [
                                    "soffice", 
                                    "--headless", 
                                    "--convert-to", "pdf", 
                                    "--outdir", os.path.dirname(temp_file_path),
                                    temp_file_path
                                ]
                                
                                result = subprocess.run(cmd, capture_output=True, text=True)
                                
                                if result.returncode == 0:
                                    pdf_output_path = os.path.splitext(temp_file_path)[0] + ".pdf"
                                    logger.info(f"Successfully converted DOCX to PDF using LibreOffice")
                                else:
                                    raise Exception(f"LibreOffice conversion failed: {result.stderr}")
                                    
                            except subprocess.CalledProcessError:
                                # LibreOffice not available, try alternative method
                                logger.warning("LibreOffice not available, trying alternative conversion method")
                                
                                # Try using python-docx-pdf if installed
                                try:
                                    from docx2pdf import convert
                                    convert(temp_file_path, pdf_output_path)
                                    logger.info(f"Successfully converted DOCX to PDF using docx2pdf")
                                except ImportError:
                                    raise Exception("No DOCX to PDF conversion tool available")
                            
                            # Now extract text from the PDF, which Gemini does support
                            with open(pdf_output_path, 'rb') as f:
                                pdf_data = f.read()
                            
                            prompt = """You are a professional multilanguage translator with a deep knowledge of HTML. Analyze this document and extract its content with precise structural preservation, extracting the content and formatting it in HTML:

1. Content Organization:
   - Maintain the original hierarchical structure (headers, sections, subsections)
   - IMPORTANT: In cases where the structure is messy, or you can't understand the structure of analyzed document, or if the document is unstructured, make sure to add some structure at your discretion to make the text readable.
   - IMPORTANT: NEVER GENERATE HTML FOR IMAGES. ALWAYS SKIP IMAGES. IF there is an image inside the document, JUST STKIP IT. Process text only, and it's formatting. The Output Must never have any <img. tags, if the image without any text is identified, skip it. If around the image tehre's a text, translate text only.
   - Preserve paragraph boundaries and logical content grouping
   - Maintain chronological or numerical sequence where present
   - Take special attention to tables, if there are any. Sometimes 1 row/column can include several rows/columns insidet them, so preseve the exact formatting how it's in the document. MAKE SURE TO ALWAYS CREATE BORDERS BETWEEN CELLS WHEN YOU CREATE TABLES. Just simple tables without any complex styling.
   - If the text is splitted to columns, but there are no borders between the columns, add some borders (full table).
   - DO NOT Include pages count. 
   - Make sure to format lists properly. Each bullet (numbered or not), should be on separate string. Bullets must be simple regardless of how they are presented in the document. Just simple bullets.

2. Formatting Guidelines:
   - Maintain all the styles, including bolden, italic or other types of formatting.
   - Preserve tabular data relationships
   - Maintain proper indentation to show hierarchical relationships
   - Keep contextually related numbers, measurements, or values together with their labels

3. Special Handling:
   - For lists of measurements/values, keep all parameters and their values together
   - For date-based content, ensure dates are formatted consistently as section headers
   - For forms or structured data, preserve the relationship between fields and values
   - For technical/scientific data, maintain the relationship between identifiers and their measurements
   - If it is an instruction/technical documentation/manual with images, make sure to translate text and preserve all the text that will be around images of the object - just create a list for this case.

4. Layout Preservation:
   - Identify when content is presented in columns and preserve column relationships
   - Maintain spacing that indicates logical grouping in the original
   - Preserve the flow of information in a way that maintains readability

5. HTML Considerations:
   - Properly handle tables by maintaining row and column relationships
   - When converting to HTML, use semantic tags to represent the document structure (<h1>, <p>, <ul>, <table>, etc.)
   - Ensure any HTML output is valid and properly nested
   - Make sure text has paragraphs and they are not together, but logically splited with <p> and <br> tags so the text is readable. 
   
Extract the content so it looks like in the initial document as much as possible. The result should be clean, structured text that accurately represents the original document's organization and information hierarchy."""
                            
                            # Use Gemini with the PDF, which is a supported format
                            response = translation_service.extraction_model.generate_content(
                                contents=[
                                    prompt,
                                    {
                                        "mime_type": "application/pdf", 
                                        "data": pdf_data
                                    }
                                ],
                                generation_config={"temperature": 0.1}
                            )
                            
                            # Clean up the temporary PDF
                            if os.path.exists(pdf_output_path):
                                os.remove(pdf_output_path)
                                
                        except Exception as pdf_conversion_error:
                            logger.error(f"PDF conversion failed: {str(pdf_conversion_error)}")
                            
                            # Last resort: Extract just the text and send as plain text to Gemini
                            import docx
                            doc = docx.Document(temp_file_path)
                            plain_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                            
                            # For tables, extract and add with clear markers
                            for table in doc.tables:
                                plain_text += "\n\n--- TABLE START ---\n"
                                for i, row in enumerate(table.rows):
                                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                                    plain_text += row_text + "\n"
                                plain_text += "--- TABLE END ---\n\n"
                            
                            prompt = """You are a professional multilanguage translator with a deep knowledge of HTML. Analyze this document and extract its content with precise structural preservation, extracting the content and formatting it in HTML:

1. Content Organization:
   - Maintain the original hierarchical structure (headers, sections, subsections)
   - IMPORTANT: In cases where the structure is messy, or you can't understand the structure of analyzed document, or if the document is unstructured, make sure to add some structure at your discretion to make the text readable.
   - IMPORTANT: NEVER GENERATE HTML FOR IMAGES. ALWAYS SKIP IMAGES. IF there is an image inside the document, JUST STKIP IT. Process text only, and it's formatting. The Output Must never have any <img. tags, if the image without any text is identified, skip it. If around the image tehre's a text, translate text only.
   - Preserve paragraph boundaries and logical content grouping
   - Maintain chronological or numerical sequence where present
   - Take special attention to tables, if there are any. Sometimes 1 row/column can include several rows/columns insidet them, so preseve the exact formatting how it's in the document. MAKE SURE TO ALWAYS CREATE BORDERS BETWEEN CELLS WHEN YOU CREATE TABLES. Just simple tables without any complex styling.
   - If the text is splitted to columns, but there are no borders between the columns, add some borders (full table).
   - DO NOT Include pages count. 
   - Make sure to format lists properly. Each bullet (numbered or not), should be on separate string. Bullets must be simple regardless of how they are presented in the document. Just simple bullets.

2. Formatting Guidelines:
   - Maintain all the styles, including bolden, italic or other types of formatting.
   - Preserve tabular data relationships
   - Maintain proper indentation to show hierarchical relationships
   - Keep contextually related numbers, measurements, or values together with their labels

3. Special Handling:
   - For lists of measurements/values, keep all parameters and their values together
   - For date-based content, ensure dates are formatted consistently as section headers
   - For forms or structured data, preserve the relationship between fields and values
   - For technical/scientific data, maintain the relationship between identifiers and their measurements
   - If it is an instruction/technical documentation/manual with images, make sure to translate text and preserve all the text that will be around images of the object - just create a list for this case.

4. Layout Preservation:
   - Identify when content is presented in columns and preserve column relationships
   - Maintain spacing that indicates logical grouping in the original
   - Preserve the flow of information in a way that maintains readability

5. HTML Considerations:
   - Properly handle tables by maintaining row and column relationships
   - When converting to HTML, use semantic tags to represent the document structure (<h1>, <p>, <ul>, <table>, etc.)
   - Ensure any HTML output is valid and properly nested
   - Make sure text has paragraphs and they are not together, but logically splited with <p> and <br> tags so the text is readable. 
   
Extract the content so it looks like in the initial document as much as possible. The result should be clean, structured text that accurately represents the original document's organization and information hierarchy."""
                            
                            # Send the plain text to Gemini
                            response = translation_service.extraction_model.generate_content(
                                prompt + "\n\n" + plain_text,
                                generation_config={"temperature": 0.1}
                            )
                        
                        html_content = response.text.strip()
                        html_content = html_content.replace('```html', '').replace('```', '').strip()
                        
                        # Add CSS styles if needed
                        if '<style>' not in html_content:
                            css_styles = """
            <style>
                .document {
                    width: 100%;
                    max-width: 1000px;
                    margin: 0 auto;
                    font-family: Arial, sans-serif;
                    line-height: 1.5;
                }
                .text-content {
                    margin-bottom: 1em;
                }
                .form-section {
                    margin-bottom: 1em;
                }
                .form-row {
                    display: flex;
                    margin-bottom: 0.5em;
                    gap: 1em;
                }
                .label {
                    width: 200px;
                    flex-shrink: 0;
                }
                .value {
                    flex-grow: 1;
                }
                .data-table {
                    width: 100%;
                    border-collapse: collapse;
                    margin-bottom: 1em;
                }
                .data-table:not(.no-borders) td,
                .data-table:not(.no-borders) th {
                    border: 1px solid black;
                    padding: 0.5em;
                }
                .no-borders td,
                .no-borders th {
                    border: none !important;
                }
                .header {
                    text-align: right;
                    margin-bottom: 20px;
                }
            </style>
            """
                            html_content = f"{css_styles}\n{html_content}"
                        
                        # Process and normalize index numbers
                        soup = BeautifulSoup(html_content, 'html.parser')
                        for index_div in soup.find_all(class_='index'):
                            index_text = index_div.get_text().strip()
                            # Use the local normalize_index method
                            corrected_index = self.normalize_index(index_text) if hasattr(self, 'normalize_index') else index_text
                            if corrected_index != index_text:
                                index_div.string = corrected_index
                                
                        html_content = str(soup)
                        
                        logger.info(f"Successfully extracted HTML content from document using base64 encoding: {len(html_content)} chars")
                        
                        return html_content
                    
                except Exception as e:
                    logger.error(f"Error processing document: {str(e)}")
                    # Create a basic fallback HTML content
                    html_content = f"""
                    <div class='document'>
                        <h1>Document Content</h1>
                        <p>The system encountered an error extracting content from your document. Here's what we know:</p>
                        <p>File type: {file_type}</p>
                        <p>Error details: {str(e)}</p>
                        <p>Please try converting your document to PDF format for better results.</p>
                    </div>
                    """
                    return html_content
                    
                finally:
                    # Clean up the temporary file
                    try:
                        os.remove(temp_file_path)
                    except Exception as e:
                        logger.warning(f"Could not delete temp file {temp_file_path}: {e}")
                    
            else:
                # Unsupported text document type
                return f"<div class='text-content'>Unsupported document type: {file_type}</div>"
                
        except Exception as e:
            logger.error(f"Error processing text document: {str(e)}")
            return f"<div class='error'>Error processing document: {str(e)}</div>"

    # Helper method to normalize index numbers
    def normalize_index(self, index_text):
        """Normalize index numbers by fixing common OCR and formatting errors"""
        if not index_text:
            return index_text
            
        # Fix common OCR errors in index numbers
        index_text = index_text.strip()
        index_text = re.sub(r'[Ll]\.', '1.', index_text)  # Replace 'L.' with '1.'
        
        # Replace incorrect decimal separators
        index_text = re.sub(r'(\d+)[,;](\d+)', r'\1.\2', index_text)
        
        # Fix missing dots
        if re.match(r'^\d+$', index_text):
            index_text = index_text + '.'
            
        # Fix spacing issues
        index_text = re.sub(r'\s+', '', index_text)
        
        # Fix incorrect indices
        index_text = re.sub(r'1\.1\.141', '1.1.1.4.1', index_text)
        index_text = re.sub(r'1\.1\.1\.42', '1.1.1.4.2', index_text)

        return index_text
    
document_processing_service = DocumentProcessingService()